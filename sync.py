# === Fixed Unicode Handling in sync.py ===
import os
import sqlite3
import hashlib
import json
import uuid
import sys
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import logging
import pandas as pd
import sys

# === FIX 1: Set UTF-8 encoding for stdout/stderr ===
if sys.platform.startswith('win'):
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.detach())
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.detach())

import os
import sys

# Detect if running in Docker container
def is_running_in_docker():
    return os.path.exists('/.dockerenv') or os.environ.get('DOCKER_CONTAINER') == 'true'

# Set KB_PATH based on environment
if is_running_in_docker():
    KB_PATH = "/app/chatbotKB_test"
else:
    # Use the actual Windows path when running locally
    KB_PATH = r"C:\Users\maria selciya\Desktop\chatbotKB_test"

# Check if the path exists
if not os.path.exists(KB_PATH):
    print(f"ERROR: KB_PATH does not exist: {KB_PATH}")
    print("Please create the directory or update KB_PATH in the script")
    sys.exit(1)

print(f"Using KB_PATH: {KB_PATH}")

# Rest of your sync.py code here...

# === FIX 2: Safe print function that handles Unicode ===
def safe_print(message):
    """Safely print messages, handling Unicode encoding issues"""
    try:
        print(message)
    except UnicodeEncodeError:
        # Fallback: print without emojis/special chars
        safe_message = message.encode('ascii', errors='ignore').decode('ascii')
        print(safe_message)
    except Exception as e:
        print(f"Print error: {str(e)}")

# Qdrant imports - using the same pattern as your working debug script
try:
    from qdrant_client import QdrantClient
    from langchain_qdrant import QdrantVectorStore
    from qdrant_client.models import Distance, VectorParams, Filter, FieldCondition, MatchValue
    QDRANT_AVAILABLE = True
except ImportError:
    safe_print("ERROR: Qdrant imports failed. Please install: pip install qdrant-client langchain-qdrant")
    QDRANT_AVAILABLE = False

# DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Logging setup
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration
load_dotenv()

#KB_PATH = os.getenv("KB_PATH", "./chatbotKB_test")
SQLITE_PATH = "./file_index.db"
COLLECTION_NAME = "kb_collection"
SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.csv', '.docx', '.md', '.py', '.js', '.html', '.xml', '.json'}

# Validate KB_PATH exists
if not os.path.exists(KB_PATH):
    safe_print(f"ERROR: KB_PATH does not exist: {KB_PATH}")
    safe_print(f"Please create the directory or update KB_PATH in the script")
    exit(1)

class KnowledgeBaseSync:
    def __init__(self):
        self.embedding_model = None
        self.vectorstore = None
        self.client = None
        self.conn = None
        self.cursor = None
        
    def initialize(self):
        """Initialize all components"""
        load_dotenv()
        
        # Check required environment variables
        required_vars = ['QDRANT_URL', 'QDRANT_API_KEY', 'GOOGLE_API_KEY']
        missing_vars = []
        
        for var in required_vars:
            if not os.getenv(var):
                missing_vars.append(var)
        
        if missing_vars:
            raise ValueError(f"Missing environment variables: {missing_vars}")
        
        google_api_key = os.getenv("GOOGLE_API_KEY")
        qdrant_url = os.getenv("QDRANT_URL")
        qdrant_api_key = os.getenv("QDRANT_API_KEY")
        
        safe_print(f"Connecting to Qdrant Cloud: {qdrant_url}")
        
        # Initialize Qdrant CLOUD client - same as your working debug script
        self.client = QdrantClient(
            url=qdrant_url,
            api_key=qdrant_api_key,
        )
        
        # Initialize embeddings - same as your working debug script
        self.embedding_model = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        
        # Test connection and get/create collection
        try:
            collections = self.client.get_collections()
            collection_names = [col.name for col in collections.collections]
            
            if COLLECTION_NAME in collection_names:
                collection_info = self.client.get_collection(COLLECTION_NAME)
                logger.info(f"Connected to existing collection '{COLLECTION_NAME}' - Points: {collection_info.points_count}")
            else:
                logger.info(f"Creating new collection: {COLLECTION_NAME}")
                self.client.create_collection(
                    collection_name=COLLECTION_NAME,
                    vectors_config=VectorParams(size=768, distance=Distance.COSINE)
                )
                logger.info("Collection created successfully!")
                
        except Exception as e:
            logger.error(f"Qdrant connection failed: {e}")
            raise
        
        # Initialize vectorstore - same pattern as debug script
        self.vectorstore = QdrantVectorStore(
            client=self.client,
            collection_name=COLLECTION_NAME,
            embedding=self.embedding_model
        )
        
        # Initialize database with migration
        self.init_database()
        logger.info("All components initialized for CLOUD sync")
    
    def check_column_exists(self, table_name, column_name):
        """Check if a column exists in a table"""
        self.cursor.execute(f"PRAGMA table_info({table_name})")
        columns = [column[1] for column in self.cursor.fetchall()]
        return column_name in columns
    
    def migrate_database_schema(self):
        """Migrate database schema to latest version"""
        logger.info("Checking database schema...")
        
        # Check if table exists
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='file_index'")
        table_exists = self.cursor.fetchone() is not None
        
        if not table_exists:
            logger.info("Creating new file_index table...")
            self.cursor.execute("""
                CREATE TABLE file_index (
                    filename TEXT PRIMARY KEY,
                    filepath TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size INTEGER,
                    mtime REAL NOT NULL,
                    hash TEXT NOT NULL,
                    chunk_count INTEGER DEFAULT 0,
                    indexed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
        else:
            # Check and add missing columns
            columns_to_add = [
                ('filepath', 'TEXT NOT NULL DEFAULT ""'),
                ('file_size', 'INTEGER DEFAULT 0'),
                ('chunk_count', 'INTEGER DEFAULT 0'),
                ('indexed_at', 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP')
            ]
            
            for column_name, column_def in columns_to_add:
                if not self.check_column_exists('file_index', column_name):
                    logger.info(f"Adding missing column: {column_name}")
                    self.cursor.execute(f"ALTER TABLE file_index ADD COLUMN {column_name} {column_def}")
            
            # Update filepath for existing records if empty
            self.cursor.execute("UPDATE file_index SET filepath = filename WHERE filepath = '' OR filepath IS NULL")
        
        self.conn.commit()
        logger.info("Database schema migration completed")
    
    def init_database(self):
        """Initialize SQLite database with migration"""
        self.conn = sqlite3.connect(SQLITE_PATH, timeout=30)
        self.cursor = self.conn.cursor()
        
        # Perform schema migration
        self.migrate_database_schema()
    
    def get_file_hash(self, file_path):
        """Calculate file hash"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.sha256(f.read()).hexdigest()
        except Exception as e:
            logger.error(f"Failed to hash {file_path}: {e}")
            return None
    
    def is_supported(self, file_path):
        """Check if file is supported"""
        return file_path.suffix.lower() in SUPPORTED_EXTENSIONS
    
    def load_document(self, file_path):
        """Load document based on file type"""
        docs = []
        ext = file_path.suffix.lower()
        
        try:
            if ext == '.pdf':
                loader = PyPDFLoader(str(file_path))
                docs = loader.load()
                docs = [doc for doc in docs if doc.page_content.strip()]
            
            elif ext == '.docx' and DOCX_AVAILABLE:
                doc = DocxDocument(str(file_path))
                content = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                if content.strip():
                    docs = [Document(page_content=content, metadata={'source': str(file_path)})]
            
            elif ext == '.csv':
                try:
                    df = pd.read_csv(file_path)
                    if not df.empty:
                        content = f"CSV: {file_path.name}\nColumns: {', '.join(df.columns)}\nRows: {len(df)}\n\n{df.to_string(max_rows=100)}"
                        docs = [Document(page_content=content, metadata={'source': str(file_path)})]
                except Exception as e:
                    logger.error(f"Failed to load CSV {file_path}: {e}")
                    return []
            
            elif ext in ['.txt', '.md', '.py', '.js', '.html', '.xml', '.json']:
                encodings = ['utf-8', 'latin-1', 'cp1252']
                content = None
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                
                if content and content.strip():
                    docs = [Document(page_content=content, metadata={'source': str(file_path)})]
            
            # Add metadata
            for doc in docs:
                doc.metadata.update({
                    'filename': file_path.name,
                    'file_type': ext,
                    'processed_at': datetime.now().isoformat()
                })
            
            return [doc for doc in docs if doc.page_content.strip()]
            
        except Exception as e:
            logger.error(f"Failed to load {file_path.name}: {e}")
            return []
    
    def remove_from_vectorstore(self, filename):
        """Remove all vectors for a specific file using Qdrant client directly"""
        try:
            # Use scroll to find points with the filename
            scroll_result = self.client.scroll(
                collection_name=COLLECTION_NAME,
                scroll_filter=Filter(
                    must=[FieldCondition(key="metadata.filename", match=MatchValue(value=filename))]
                ),
                limit=10000  # Get all points for this file
            )
            
            point_ids = [point.id for point in scroll_result[0]]
            
            if point_ids:
                self.client.delete(
                    collection_name=COLLECTION_NAME,
                    points_selector=point_ids
                )
                logger.info(f"Removed {len(point_ids)} vectors for: {filename}")
            else:
                logger.info(f"No vectors found for: {filename}")
            
            return True
        except Exception as e:
            logger.error(f"Failed to remove vectors for {filename}: {e}")
            return False
    
    def add_to_vectorstore(self, file_path):
        """Add file to vector store"""
        try:
            filename = file_path.name
            safe_print(f"PROCESSING: {filename}")
            
            # Remove existing vectors first
            self.remove_from_vectorstore(filename)
            
            # Load documents
            docs = self.load_document(file_path)
            if not docs:
                logger.warning(f"No content to index for {filename}")
                return False
            
            # Split into chunks
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = []
            for doc in docs:
                doc_chunks = splitter.split_documents([doc])
                chunks.extend(doc_chunks)
            
            if not chunks:
                logger.warning(f"No chunks created for {filename}")
                return False
            
            # Add filename and chunk_id to all chunks
            for i, chunk in enumerate(chunks):
                chunk.metadata['filename'] = filename
                chunk.metadata['chunk_id'] = str(uuid.uuid4())
                chunk.metadata['chunk_index'] = i
            
            # Add to vectorstore in batches to avoid timeout
            batch_size = 50
            total_added = 0
            
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                try:
                    ids = self.vectorstore.add_documents(batch)
                    total_added += len(batch)
                    safe_print(f"   Batch {i//batch_size + 1}: Added {len(batch)} chunks")
                except Exception as e:
                    logger.error(f"Failed to add batch {i//batch_size + 1}: {e}")
                    # Continue with next batch
            
            if total_added == 0:
                return False
            
            # Update database
            file_stats = file_path.stat()
            file_hash = self.get_file_hash(file_path)
            
            self.cursor.execute("""
                INSERT OR REPLACE INTO file_index 
                (filename, filepath, file_type, file_size, mtime, hash, chunk_count, indexed_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            """, (filename, str(file_path), file_path.suffix.lower(), 
                  file_stats.st_size, file_stats.st_mtime, file_hash, total_added))
            self.conn.commit()
            
            safe_print(f"ADDED: {filename} ({total_added} chunks)")
            return True
            
        except Exception as e:
            logger.error(f"Failed to add {file_path.name}: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def remove_file(self, filename):
        """Completely remove file from database and vectorstore"""
        try:
            # Remove from vectorstore
            self.remove_from_vectorstore(filename)
            
            # Remove from database
            self.cursor.execute("DELETE FROM file_index WHERE filename = ?", (filename,))
            self.conn.commit()
            
            logger.info(f"Completely removed: {filename}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to remove {filename}: {e}")
            return False
    
    def get_current_files(self):
        """Get current files from filesystem"""
        kb_dir = Path(KB_PATH)
        if not kb_dir.exists():
            raise FileNotFoundError(f"KB directory not found: {KB_PATH}")
        
        current_files = {}
        for file_path in kb_dir.rglob("*"):
            if (file_path.is_file() and 
                not file_path.name.startswith('.') and 
                self.is_supported(file_path)):
                
                file_hash = self.get_file_hash(file_path)
                if file_hash:
                    current_files[file_path.name] = {
                        'path': file_path,
                        'mtime': file_path.stat().st_mtime,
                        'hash': file_hash
                    }
        
        return current_files
    
    def get_database_files(self):
        """Get files from database"""
        self.cursor.execute("SELECT filename, mtime, hash FROM file_index")
        return {row[0]: {'mtime': row[1], 'hash': row[2]} for row in self.cursor.fetchall()}
    
    def print_current_files(self):
        """Print current files in database"""
        self.cursor.execute("""
            SELECT filename, file_type, chunk_count, indexed_at 
            FROM file_index 
            ORDER BY indexed_at DESC
        """)
        files = self.cursor.fetchall()
        
        safe_print("\n" + "="*60)
        safe_print("CURRENT DATABASE FILES:")
        safe_print("="*60)
        
        if files:
            for filename, file_type, chunk_count, indexed_at in files:
                safe_print(f"SUCCESS: {filename} ({file_type}) - {chunk_count} chunks")
        else:
            safe_print("No files in database")
        
        safe_print(f"Total: {len(files)} files")
        safe_print("="*60)
    
    def rebuild_database(self):
        """Rebuild database from scratch"""
        logger.info("Rebuilding database from scratch...")
        
        # Drop and recreate table
        self.cursor.execute("DROP TABLE IF EXISTS file_index")
        self.cursor.execute("""
            CREATE TABLE file_index (
                filename TEXT PRIMARY KEY,
                filepath TEXT NOT NULL,
                file_type TEXT NOT NULL,
                file_size INTEGER DEFAULT 0,
                mtime REAL NOT NULL,
                hash TEXT NOT NULL,
                chunk_count INTEGER DEFAULT 0,
                indexed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        self.conn.commit()
        logger.info("Database rebuilt")
    
    def test_search(self):
        """Test search functionality"""
        try:
            safe_print("\nTesting search functionality...")
            results = self.vectorstore.similarity_search("test query", k=1)
            if results:
                safe_print(f"Search working! Found {len(results)} results")
                safe_print(f"Sample: {results[0].page_content[:100]}...")
            else:
                safe_print("No results found (collection might be empty)")
            return True
        except Exception as e:
            safe_print(f"Search test failed: {e}")
            return False
    
    def sync(self, force_rebuild=False):
        """Main sync function - optimized for real-time updates"""
        try:
            logger.info("Starting knowledge base sync...")
            
            if not QDRANT_AVAILABLE:
                safe_print("Qdrant not available. Please install required packages.")
                return False
            
            if not self.client:
                self.initialize()
            
            # If force rebuild, rebuild database
            if force_rebuild:
                self.rebuild_database()
            
            # Get current state
            current_files = self.get_current_files()
            db_files = self.get_database_files()
            
            safe_print(f"Found {len(current_files)} supported files in directory")
            
            if len(current_files) == 0:
                safe_print(f"WARNING: No supported files found in {KB_PATH}")
                safe_print(f"Supported extensions: {SUPPORTED_EXTENSIONS}")
                return False
            
            # Find changes
            to_add = []
            to_remove = []
            
            # Check for new/modified files
            for filename, file_info in current_files.items():
                if (force_rebuild or 
                    filename not in db_files or 
                    abs(db_files[filename]['mtime'] - file_info['mtime']) > 1 or
                    db_files[filename]['hash'] != file_info['hash']):
                    to_add.append(file_info['path'])
            
            # Check for deleted files
            for filename in db_files:
                if filename not in current_files:
                    to_remove.append(filename)
            
            # Process changes
            added_count = 0
            failed_count = 0
            
            # Remove deleted files
            for filename in to_remove:
                if self.remove_file(filename):
                    safe_print(f"REMOVED: {filename}")
                else:
                    safe_print(f"FAILED TO REMOVE: {filename}")
            
            # Add/update files
            for file_path in to_add:
                if self.add_to_vectorstore(file_path):
                    added_count += 1
                else:
                    failed_count += 1
                    safe_print(f"FAILED: {file_path.name}")
            
            # Test search if we have data
            if added_count > 0 or len(db_files) > 0:
                self.test_search()
            
            # Show final state
            self.print_current_files()
            
            # Summary
            safe_print(f"\nSYNC COMPLETE!")
            safe_print(f"Added/Updated: {added_count}")
            safe_print(f"Removed: {len(to_remove)}")
            safe_print(f"Failed: {failed_count}")
            
            if added_count == 0 and len(to_remove) == 0:
                safe_print("No changes detected - database is up to date")
            
            return True
            
        except Exception as e:
            logger.error(f"Sync failed: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def cleanup(self):
        """Clean up resources"""
        if self.conn:
            self.conn.close()

import os
import sys
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def is_running_in_docker():
    """Check if script is running inside Docker container"""
    return os.path.exists('/.dockerenv') or os.environ.get('DOCKER_CONTAINER') == 'true'

def get_kb_path():
    """Get the correct KB path based on environment"""
    if is_running_in_docker():
        # Inside Docker container, use the mounted path
        kb_path = "/app/chatbotKB_test"
    else:
        # Outside Docker, use the local Windows path from .env
        kb_path = os.getenv('KB_PATH', r'C:\Users\maria selciya\Desktop\chatbotKB_test')
        # Convert forward slashes to backslashes for Windows
        if os.name == 'nt':  # Windows
            kb_path = kb_path.replace('/', '\\')
    
    return kb_path

# Get the appropriate KB path
KB_PATH = get_kb_path()

print(f"Environment: {'Docker' if is_running_in_docker() else 'Local'}")
print(f"Using KB_PATH: {KB_PATH}")

# Check if the path exists
if not os.path.exists(KB_PATH):
    error_msg = f"ERROR: KB_PATH does not exist: {KB_PATH}"
    print(error_msg)
    
    if is_running_in_docker():
        print("Please ensure the volume is properly mounted in docker-compose.yml")
    else:
        print("Please create the directory or update KB_PATH in your .env file")
    
    sys.exit(1)

print("KB_PATH exists, proceeding with sync...")

# Your existing sync logic here
try:
    # Add your actual sync code here
    print("Sync completed successfully!")
    
except Exception as e:
    print(f"Sync failed: {str(e)}")
    sys.exit(1)
    
# Main execution functions
def sync_kb(force_rebuild=False):
    """Convenient function to sync knowledge base"""
    syncer = KnowledgeBaseSync()
    try:
        return syncer.sync(force_rebuild)
    finally:
        syncer.cleanup()

def list_files():
    """List current files in database"""
    syncer = KnowledgeBaseSync()
    try:
        syncer.initialize()
        syncer.print_current_files()
    finally:
        syncer.cleanup()

def test_connection():
    """Test Qdrant cloud connection"""
    syncer = KnowledgeBaseSync()
    try:
        syncer.initialize()
        safe_print("Connection test successful!")
        syncer.test_search()
        return True
    except Exception as e:
        safe_print(f"Connection test failed: {e}")
        return False
    finally:
        syncer.cleanup()

# CLI interface
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == "--rebuild":
        safe_print("Running with --rebuild flag (will recreate database)")
        sync_kb(force_rebuild=True)
    elif len(sys.argv) > 1 and sys.argv[1] == "--list":
        safe_print("Listing current files in database")
        list_files()
    elif len(sys.argv) > 1 and sys.argv[1] == "--test":
        safe_print("Testing connection")
        test_connection()
    else:
        safe_print("Running normal sync")
        sync_kb()